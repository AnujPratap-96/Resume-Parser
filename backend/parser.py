import json
import logging
import os
import re
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from dotenv import load_dotenv, find_dotenv
from groq import Groq
from pypdf import PdfReader
from docx import Document

from models import (
    JobDescription,
    Resume,
    MatchResult,
    AnalysisResponse,
)

load_dotenv(find_dotenv())

logger = logging.getLogger(__name__)

MODEL = os.getenv("LLM_MODEL", "llama-3.3-70b-versatile")
LLM_TIMEOUT = float(os.getenv("LLM_TIMEOUT", "20"))
LLM_MAX_RETRIES = max(1, int(os.getenv("LLM_MAX_RETRIES", "2")))

# Weights used to recombine score_breakdown into overall_score.
SCORE_WEIGHTS = {"skills": 0.6, "experience": 0.3, "education": 0.1}


class ConfigError(RuntimeError):
    """Server is misconfigured (e.g. no API key) — not the caller's fault."""


class ResumeReadError(ValueError):
    """Uploaded file could not be turned into text — safe to show the user."""


_client: Groq | None = None
_client_lock = threading.Lock()


def has_api_key() -> bool:
    return bool(os.getenv("GROQ_API_KEY"))


def _get_client() -> Groq:
    """Build the Groq client lazily.

    Built at first use rather than at import so the module can be imported
    (and tested) without an API key in the environment.
    """
    global _client
    if _client is None:
        with _client_lock:
            if _client is None:
                key = os.getenv("GROQ_API_KEY")
                if not key:
                    raise ConfigError("GROQ_API_KEY is not set")
                _client = Groq(api_key=key, timeout=LLM_TIMEOUT)
    return _client


# ── File readers ──────────────────────────────────────────────

CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

SCANNED_PDF_MESSAGE = (
    "Could not extract text from this PDF — it is likely a scanned "
    "(image-based) document. Resumes with photos/emoji as images need "
    "OCR. Convert it to DOCX and try again."
)


def clean_text(text: str) -> str:
    """Remove control characters and normalize whitespace.

    Handles PDFs that contain emoji, special glyphs or stray bytes —
    keeps meaningful Unicode (including emoji) but strips junk.
    """
    text = CONTROL_CHARS.sub("", text or "")
    text = text.replace("\u200b", "").replace("\ufeff", "")  # zero-width + BOM
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _ocr_pdf(file_path: str | Path) -> str:
    """Last-resort OCR for scanned PDFs.

    Needs pdf2image + pytesseract AND the poppler/tesseract binaries. A missing
    binary raises OSError (TesseractNotFoundError), not ImportError — catching
    only ImportError let that escape as a 500.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract

        parts = [
            pytesseract.image_to_string(img)
            for img in convert_from_path(str(file_path), dpi=200)
        ]
        cleaned = clean_text("\n".join(parts))
    except Exception as exc:  # missing package, missing binary, or OCR failure
        logger.warning("OCR fallback unavailable: %s", exc)
        raise ResumeReadError(SCANNED_PDF_MESSAGE) from exc

    if len(cleaned) < 50:
        raise ResumeReadError(
            "Text extraction produced too little content. The PDF may be "
            "scanned/image-only or corrupted."
        )
    return cleaned


def read_pdf(file_path: str | Path) -> str:
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if not text or len(text.strip()) < 20:
            text = page.extract_text(extraction_mode="layout")
        if text and len(text.strip()) >= 20:
            pages.append(text)
    cleaned = clean_text("\n".join(pages))
    if len(cleaned) < 50:
        cleaned = _ocr_pdf(file_path)
    return cleaned


def read_docx(file_path: str | Path) -> str:
    doc = Document(file_path)
    lines = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                if paragraph.text.strip():
                    lines.append(paragraph.text)
    return clean_text("\n".join(lines))


def read_resume(file_path: str | Path) -> str:
    p = Path(file_path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(p)
    if suffix == ".docx":
        return read_docx(p)
    raise ResumeReadError(f"Unsupported file type '{suffix}'. Upload a PDF or DOCX.")


# ── LLM helpers ──────────────────────────────────────────────

def _llm_json(messages: list[dict], max_retries: int = LLM_MAX_RETRIES) -> dict:
    attempts = max(1, max_retries)
    last_error: Exception | None = None
    for attempt in range(attempts):
        try:
            resp = _get_client().chat.completions.create(
                model=MODEL,
                messages=messages,
                response_format={"type": "json_object"},
            )
            content = clean_text(resp.choices[0].message.content or "")
            return json.loads(content)
        except ConfigError:
            raise  # retrying a missing API key is pointless
        except Exception as exc:
            last_error = exc
            if attempt < attempts - 1:
                time.sleep(2 ** attempt)
    raise last_error if last_error else RuntimeError("LLM call failed")


# ── JD parser ────────────────────────────────────────────────

def parse_job_description(text: str) -> JobDescription:
    schema = JobDescription.model_json_schema()
    messages = [
        {
            "role": "system",
            "content": (
                "You are an expert HR assistant. Extract structured "
                "information from the job description.\n\n"
                f"Return ONLY valid JSON matching this schema:\n{schema}\n\n"
                "Rules:\n"
                "- Do not return the schema itself.\n"
                "- Fill with actual data from the JD.\n"
                "- If minimum experience is missing, return null.\n"
                "- If a list is empty, return [].\n"
                "- Do not invent information."
            ),
        },
        {"role": "user", "content": f"Analyze this job description:\n\n{text}"},
    ]
    data = _llm_json(messages)
    return JobDescription(**data)


# ── Resume parser ────────────────────────────────────────────

def parse_resume(text: str) -> Resume:
    schema = Resume.model_json_schema()
    messages = [
        {
            "role": "system",
            "content": (
                "You are an expert resume parser.\n\n"
                "Extract information based on meaning, not section headings.\n"
                "Resumes may use different headings for the same content.\n"
                "Include internships under experiences.\n"
                "Extract skills mentioned anywhere in the resume.\n"
                "Extract GitHub and LinkedIn profile URLs if present.\n"
                "For each experience, extract key achievement highlights as a list.\n\n"
                f"Return ONLY valid JSON matching this schema:\n{schema}\n\n"
                "Rules:\n"
                "- Do not invent information.\n"
                "- If a value is missing, return null.\n"
                "- If a list is empty, return []."
            ),
        },
        {"role": "user", "content": f"Parse this resume:\n\n{text}"},
    ]
    data = _llm_json(messages)
    return Resume(**data)


# ── Scorer ────────────────────────────────────────────────────

def compute_match(job: JobDescription, resume: Resume, semantic_pairs: list[dict] | None = None) -> MatchResult:
    semantic_block = ""
    if semantic_pairs:
        lines = [
            f"- JD skill '{p['jd_skill']}' was matched to resume skill '{p['matched_skill']}' "
            f"({p['similarity']}% similarity) — treat these as the same skill."
            for p in semantic_pairs
        ]
        semantic_block = (
            "\nSemantic skill equivalences found (treat as already matched):\n"
            + "\n".join(lines)
            + "\n"
        )

    prompt = (
        "You are an experienced HR recruiter. Compare the candidate's "
        "resume against the job description and return a structured match report.\n\n"
        f"JOB DESCRIPTION:\n{job.model_dump_json(indent=2)}\n\n"
        f"CANDIDATE RESUME:\n{resume.model_dump_json(indent=2)}\n\n"
        f"{semantic_block}"
        f"Return ONLY valid JSON matching this schema:\n{MatchResult.model_json_schema()}\n\n"
        "CRITICAL CONSISTENCY RULES — violations are the worst kind of error:\n"
        "- Never list a skill as 'missing' if it appears in the candidate's "
        "skills list or the semantic equivalences above.\n"
        "- Never claim the resume 'does not mention' a technology that is in "
        "candidate.skills, projects, or experiences.skills_used.\n"
        "- Never mention a technology in weaknesses if it is in the resume.\n"
        "- A technology with a different name but the same meaning (e.g. "
        "Node.js vs NodeJS, REST API vs REST APIs) counts as present.\n"
        "- Related but distinct technologies are NOT the same skill "
        "(MySQL is not PostgreSQL, Java is not JavaScript, R is not React).\n\n"
        "Be thorough:\n"
        "1. overall_score: 0-100 based on weighted criteria\n"
        "2. skills.matched: skills from resume that match JD requirements\n"
        "3. skills.missing: skills from JD that resume truly lacks\n"
        "4. skills.extra: notable skills on resume not in JD\n"
        "5. experience: check if candidate meets minimum experience\n"
        "6. education_match: list of education items that satisfy requirements\n"
        "7. strengths/weaknesses: bullet-point insights grounded ONLY in resume data\n"
        "8. verdict: short final recommendation\n"
        "9. score_breakdown: break overall_score into individual scores (0-100 each) for skills, experience, and education\n"
        "10. improvement_tips: list of actionable suggestions to improve the resume for this role, each with area, suggestion, and impact (high/medium/low)"
    )
    messages = [{"role": "user", "content": prompt}]
    data = _llm_json(messages)
    return MatchResult(**data)


# ── High-level pipeline ──────────────────────────────────────

def analyze_full(jd_text: str, resume_path: str | Path) -> AnalysisResponse:
    """Run the full pipeline: JD parse → resume parse → match + ATS + semantic."""

    from ats import analyze_ats
    from semantic import semantic_match

    resume_text = read_resume(resume_path)

    # The JD parse and the resume parse are independent LLM calls — running
    # them together cuts roughly a third off end-to-end latency.
    with ThreadPoolExecutor(max_workers=2) as pool:
        job_future = pool.submit(parse_job_description, jd_text)
        resume_future = pool.submit(parse_resume, resume_text)
        job = job_future.result()
        resume = resume_future.result()

    semantic = semantic_match(job, resume)
    semantic_pairs = [p.model_dump() for p in semantic.pairs]
    match = compute_match(job, resume, semantic_pairs)
    match = enforce_consistency(job, resume, match, semantic_pairs)

    ats = analyze_ats(job, resume, jd_text, resume_text)

    return AnalysisResponse(
        job=job,
        resume=resume,
        match=match,
        ats=ats,
        semantic=semantic,
    )


# Word boundary that treats '#', '+' and '.' as part of the term, so "C" does
# not match inside "C++" and "node" does not match inside "node.js".
_BOUNDARY = r"(?<![\w#+.]){term}(?![\w#+.])"


def _mentions(text: str, skill: str) -> bool:
    """True if `text` refers to `skill` as a whole term."""
    if not text or not skill:
        return False
    pattern = _BOUNDARY.replace("{term}", re.escape(skill))
    return re.search(pattern, text, re.IGNORECASE) is not None


def enforce_consistency(
    job: JobDescription,
    resume: Resume,
    match: MatchResult,
    semantic_pairs: list[dict],
) -> MatchResult:
    """Deterministically fix LLM contradictions.

    A JD skill that equals a resume skill after normalization — or that the
    semantic matcher paired — is present by definition, so the LLM may not list
    it as missing or as a weakness.

    Comparison is exact on normalized terms. The old substring test
    ("jd_term in r or r in jd_term") reported "R" as covered by "React" and
    "C" as covered by "C++".
    """

    from semantic import normalize_skill

    resume_terms = {normalize_skill(s) for s in resume.skills if s}
    for exp in resume.experiences:
        resume_terms.update(normalize_skill(s) for s in exp.skills_used if s)
    for pair in semantic_pairs:
        resume_terms.add(normalize_skill(pair["matched_skill"]))
        resume_terms.add(normalize_skill(pair["jd_skill"]))
    resume_terms.discard("")

    jd_skills = [s for s in job.required_skills + job.preferred_skills if s]
    present_terms = {normalize_skill(s) for s in jd_skills} & resume_terms
    present_skills = [s for s in jd_skills if normalize_skill(s) in present_terms]

    matched = list(match.skills.matched)
    matched_terms = {normalize_skill(s) for s in matched}
    newly_matched = [
        s for s in present_skills if normalize_skill(s) not in matched_terms
    ]
    matched.extend(newly_matched)

    match.skills.matched = matched
    match.skills.missing = [
        s for s in match.skills.missing if normalize_skill(s) not in present_terms
    ]

    for skill in present_skills:
        match.weaknesses = [w for w in match.weaknesses if not _mentions(w, skill)]
        match.improvement_tips = [
            t for t in match.improvement_tips if not _mentions(t.suggestion, skill)
        ]

    if newly_matched:
        _rescore(match)

    return match


def _rescore(match: MatchResult) -> None:
    """Recompute the skills sub-score and overall score after reclassification.

    Keeps overall_score consistent with the breakdown the UI renders. The old
    code nudged overall_score by a flat +10 and left score_breakdown untouched,
    so the headline number contradicted its own bars.
    """
    total = len(match.skills.matched) + len(match.skills.missing)
    if total:
        match.score_breakdown.skills = round(
            len(match.skills.matched) / total * 100, 1
        )
    b = match.score_breakdown
    weighted = (
        b.skills * SCORE_WEIGHTS["skills"]
        + b.experience * SCORE_WEIGHTS["experience"]
        + b.education * SCORE_WEIGHTS["education"]
    )
    match.overall_score = round(min(100.0, max(0.0, weighted)), 1)
